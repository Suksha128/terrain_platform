import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_centered_bold(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_centered(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'

# --- PAGE 1: TITLE ---
doc.add_paragraph('\n\n\n')
add_centered_bold("Terrain Intelligence Platform for Precision Agriculture", 18)
doc.add_paragraph('\n')
add_centered("A REPORT")
doc.add_paragraph('\n')
add_centered("submitted by")
doc.add_paragraph('\n')
add_centered_bold("[Your Name] ( [Your Reg. No] )", 14)
doc.add_paragraph('\n')
add_centered("in partial fulfilment for the award of")
doc.add_paragraph('\n')
add_centered_bold("M. Tech. (Integrated) Software Engineering", 14)
doc.add_paragraph('\n')
add_centered_bold("School of Computer Science and Engineering", 14)
add_centered_bold("Vellore Institute of Technology", 14)
add_centered("(Deemed to be University under section 3 of UGC Act, 1956)", 10)
doc.add_paragraph('\n\n')
add_centered_bold("November - 2025", 14)

doc.add_page_break()

# --- PAGE 2: DECLARATION ---
add_centered_bold("School of Computer Science and Engineering", 14)
add_centered_bold("DECLARATION", 14)
doc.add_paragraph('\n')

p = doc.add_paragraph("I hereby declare that the project entitled ")
p.add_run("“Terrain Intelligence Platform for Precision Agriculture”").bold = True
p.add_run(" submitted by me to the School of Computer Science and Engineering, Vellore Institute of Technology, Chennai Campus, Chennai 600127 in partial fulfilment of the requirements for the award of the degree of ")
p.add_run("M. Tech. (Integrated) Software Engineering").bold = True
p.add_run(" is a record of bonafide work carried out by me. I further declare that the work reported in this report has not been submitted and will not be submitted, either in part or in full, for the award of any other degree or diploma of this institute or of any other institute or university.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n\n\nSignature\n\n\n')
p = doc.add_paragraph()
p.add_run("[Your Name] ( [Your Reg. No] )").bold = True

doc.add_page_break()

# --- PAGE 3: CERTIFICATE ---
add_centered_bold("School of Computer Science and Engineering", 14)
add_centered_bold("CERTIFICATE", 14)
doc.add_paragraph('\n')

p = doc.add_paragraph("The project report entitled ")
p.add_run("“Terrain Intelligence Platform for Precision Agriculture”").bold = True
p.add_run(" is prepared and submitted by ")
p.add_run("[Your Name] (Register No: [Your Reg. No])").bold = True
p.add_run(". It has been found satisfactory in terms of scope, quality and presentation as partial fulfilment of the requirements for the award of the degree of ")
p.add_run("M. Tech. (Integrated) Software Engineering").bold = True
p.add_run(" in Vellore Institute of Technology, Chennai, India.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n\n\nExamined by:\n\n\n\n')
p = doc.add_paragraph("Examiner I                                                                                Examiner II")
p.runs[0].bold = True

doc.add_page_break()

# --- PAGE 4: INDUSTRY CERTIFICATE ---
add_centered_bold("[INSERT THE CERTIFICATE OF MERIT OBTAINED FROM THE INDUSTRY WHERE YOU DID YOUR INTERNSHIP HERE]")
doc.add_page_break()

# --- PAGE 5: ACKNOWLEDGEMENT ---
add_centered_bold("ACKNOWLEDGEMENT", 14)
doc.add_paragraph('\n')
p = doc.add_paragraph("I would like to express my deepest gratitude to all those who provided me the possibility to complete this report.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

acks = [
    "Dr. M. Premalatha, Head of the Department, M. Tech (Integrated) Software Engineering, Vellore Institute of Technology, Chennai",
    "Dr. Viswanathan V, DEAN, School of Computer Science & Engineering, Vellore Institute of Technology, Chennai",
    "Dr. NITHYANANDAM P, Associate Dean, School of Computer Science & Engineering, Vellore Institute of Technology, Chennai",
    "Dr. SUGANYA G, Associate Dean, School of Computer Science & Engineering, Vellore Institute of Technology, Chennai",
    "Dr. SWEETLIN HEMALATHA C, Associate Dean, School of Computer Science & Engineering, Vellore Institute of Technology, Chennai",
    "I also extend my sincere thanks to my parents, friends, and industry mentors for their continuous support and guidance throughout this project."
]

for ack in acks:
    doc.add_paragraph(ack, style='List Bullet')

doc.add_page_break()

# --- PAGE 6: CONTENTS ---
add_centered_bold("CONTENTS", 14)
doc.add_paragraph('\n')

contents = [
    ("Title Page", "i"),
    ("Declaration", "ii"),
    ("Certificate", "iii"),
    ("Industry certificate", "iv"),
    ("Acknowledgement", "v"),
    ("Table of contents", "vi"),
    ("List of Abbreviations", "ix"),
    ("Abstract", "x"),
    ("1. Introduction", "01"),
    ("2. System Architecture & Methodology", "02"),
    ("3. Implementation & Results", "03"),
    ("4. Conclusion", "04"),
    ("References", "05"),
    ("Appendix – I (Source Code & Setup)", "06")
]

for title, page in contents:
    p = doc.add_paragraph(f"{title}")
    p.add_run(f"\t\t\t\t\t\t\t{page}")

doc.add_page_break()

# --- PAGE 7: ABBREVIATIONS ---
add_centered_bold("LIST OF ABBREVIATIONS", 14)
doc.add_paragraph('\n')

abbrevs = [
    ("DEM", "Digital Elevation Model"),
    ("DTM", "Digital Terrain Model"),
    ("DSM", "Digital Surface Model"),
    ("NDVI", "Normalized Difference Vegetation Index"),
    ("TWI", "Topographic Wetness Index"),
    ("SCA", "Specific Catchment Area"),
    ("GEE", "Google Earth Engine"),
    ("CRS", "Coordinate Reference System"),
    ("API", "Application Programming Interface")
]

for abbr, exp in abbrevs:
    doc.add_paragraph(f"{abbr}\t\t\t{exp}")

doc.add_page_break()

# --- PAGE 8: ABSTRACT ---
add_centered_bold("ABSTRACT", 14)
doc.add_paragraph('\n')

p = doc.add_paragraph()
p.add_run("Present Scenario of the Area: ").bold = True
p.add_run("In modern agriculture, optimizing crop yield heavily relies on understanding terrain hydrology and preventing waterlogging. Currently, agronomists rely on manual field surveys or basic 30m resolution elevation models (like SRTM) to estimate where water will pool.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
p.add_run("Limitations/Pitfalls of the Present Scenario: ").bold = True
p.add_run("Standard mapping techniques and low-resolution datasets are physically insufficient for farm-level micro-topography. Basic D8 flow routing algorithms force water into artificial 45-degree grid paths, failing to capture true topological fluid dynamics. Furthermore, static elevation models do not account for real-time weather events or actual live vegetation health, leading to inaccurate drainage planning and unexpected crop loss.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
p.add_run("Problem Addressed: ").bold = True
p.add_run("There is a critical lack of automated, high-precision, multi-variable analytical tools that can fuse advanced topological physics (like D-Infinity flow accumulation) with live environmental data (rainfall forecasts and satellite NDVI) to dynamically generate reliable agricultural management zones.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
p.add_run("Proposed Solution: ").bold = True
p.add_run("This project introduces the \"Terrain Intelligence Platform,\" an enterprise-grade geospatial analytics engine. The platform utilizes a decoupled architecture, employing a high-performance FastAPI backend to process high-resolution bare-earth LiDAR and FABDEM datasets using WhiteboxTools. By upgrading from D8 to D-Infinity Specific Catchment Area (SCA) algorithms, it physically simulates accurate water flow. The system dynamically queries the Open-Meteo API for 3-day rainfall forecasts and cross-references Topographic Wetness Indexes (TWI) with Sentinel-2 NDVI satellite imagery. A machine learning engine utilizing K-Means clustering then automatically delineates the farm into highly precise, risk-assessed management zones, rendered via an interactive Streamlit and Folium dashboard.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# --- CHAPTER 1 ---
add_centered_bold("1 INTRODUCTION", 14)
doc.add_paragraph('\n')
p = doc.add_paragraph("Precision agriculture requires exact knowledge of field topography to manage irrigation, prevent waterlogging, and maximize yield. The \"Terrain Intelligence Platform\" is developed to solve the complex problem of agricultural water routing by moving beyond basic mapping and introducing physics-based hydrological modeling.\n\nThis project tackles the integration of multi-modal geospatial data—combining Digital Terrain Models (DTMs), satellite-derived vegetation indices (NDVI), and live meteorological data into a unified, automated Machine Learning pipeline. The primary objective is to allow end-users to upload raw GeoTIFF datasets and instantly receive interactive, mathematically precise risk-management zones without requiring deep GIS expertise.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_page_break()

# --- CHAPTER 2 ---
add_centered_bold("2 SYSTEM ARCHITECTURE & METHODOLOGY", 14)
doc.add_paragraph('\n')
p = doc.add_paragraph("The system is designed with a highly decoupled architecture to handle intense geospatial raster mathematics asynchronously.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading("2.1 Backend Engine (FastAPI)", level=3)
doc.add_paragraph("The core engine is built using Python and FastAPI. It leverages WhiteboxTools, Rasterio, and GeoPandas for heavy raster math. The engine processes DTMs to calculate Slope and Topographic Wetness Index (TWI). Crucially, the hydrology engine utilizes D-Infinity (D-Inf) Flow Tracing rather than legacy D8 algorithms, allowing water to fractionally disperse across exact angular contours.")

doc.add_heading("2.2 Live Data Integration", level=3)
doc.add_paragraph("The backend extracts bounding box coordinates from the uploaded geospatial rasters and automatically queries the Open-Meteo REST API. This injects the upcoming 3-Day rainfall forecasts directly into the risk assessment heuristics.")

doc.add_heading("2.3 Machine Learning Pipeline", level=3)
doc.add_paragraph("A Scikit-Learn pipeline fuses the physical hydrology data with ground-truth NDVI. K-Means clustering is utilized within hydrological watersheds to mathematically cluster micro-zones, escalating risk levels if physical sinkholes correlate with low NDVI values.")

doc.add_heading("2.4 Frontend Dashboard (Streamlit)", level=3)
doc.add_paragraph("The user interface is powered by Streamlit. It features secure session state management to handle asynchronous API calls and renders the generated GeoJSON risk zones on a live, interactive Folium map.")
doc.add_page_break()

# --- CHAPTER 3 ---
add_centered_bold("3 IMPLEMENTATION & RESULTS", 14)
doc.add_paragraph('\n')

doc.add_heading("3.1 Data Preparation", level=3)
doc.add_paragraph("High-resolution bare-earth DTMs (such as LiDAR or FABDEM) and Sentinel-2 Multispectral imagery are pre-processed to 10-meter resolution in Google Earth Engine using bicubic resampling and projected to local UTM Coordinate Reference Systems (e.g., EPSG:32644).")

doc.add_heading("3.2 Hydrological Execution", level=3)
doc.add_paragraph("Upon uploading dtm.tif and ndvi.tif, the system successfully hydro-conditions the DEM (filling artificial depressions) and computes the Specific Catchment Area. The resulting flow_direction and watershed_basin rasters accurately reflect real-world drainage paths.")

doc.add_heading("3.3 Final Output", level=3)
doc.add_paragraph("The platform successfully outputs an interactive \"Combined Risk Zone\" map. Results demonstrated that areas with high TWI and low NDVI were accurately flagged as \"Critical Risk\" zones, proving the efficacy of combining physics-based routing with satellite ground-truth data.")
doc.add_page_break()

# --- CHAPTER 4 ---
add_centered_bold("4 CONCLUSION", 14)
doc.add_paragraph('\n')
p = doc.add_paragraph("The Terrain Intelligence Platform successfully demonstrates an enterprise-grade approach to precision agriculture analytics. By upgrading to D-Infinity algorithms and fusing static terrain physics with dynamic rainfall and crop health data, the platform provides significantly higher reliability than traditional DEM analysis. This tool empowers agronomists and industrial farmers to make highly targeted, data-driven decisions regarding sub-surface drainage and crop management, ultimately saving resources and optimizing yield.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_page_break()

# --- REFERENCES ---
add_centered_bold("REFERENCES", 14)
doc.add_paragraph('\n')
refs = [
    "[1] J. Lindsay, \"WhiteboxTools: Advanced geospatial data analysis platform,\" Guelph, Ontario, 2023.",
    "[2] L. Tarboton, \"A new method for the determination of flow directions and upslope areas in grid digital elevation models,\" Water Resources Research, vol. 33, no. 2, pp. 309-319, 1997.",
    "[3] Copernicus Sentinel-2 Multispectral Instrument (MSI) Data, European Space Agency.",
    "[4] Open-Meteo REST API Documentation, 2024."
]
for ref in refs:
    doc.add_paragraph(ref)
doc.add_page_break()

# --- APPENDIX I ---
add_centered_bold("APPENDIX I", 14)
doc.add_paragraph('\n')
doc.add_heading("Core Dependency Configuration (requirements.txt)", level=3)
doc.add_paragraph("streamlit>=1.35.0\nrequests>=2.31.0\npandas>=2.2.2\nfolium>=0.16.0\nstreamlit-folium>=0.20.0\ngeopandas>=0.14.4\nfastapi>=0.111.0\nuvicorn[standard]>=0.29.0\nrasterio>=1.3.10\nwhitebox>=2.3.4\nscikit-learn>=1.5.0")

doc.add_heading("Platform Launch Commands", level=3)
doc.add_paragraph("cd backend\npython3 -m uvicorn app.main:app --port 8000 --reload\n\npython3 -m streamlit run streamlit_app.py")

doc.save('/Users/sukshas/Downloads/terrain-platform1/Project_Report.docx')
